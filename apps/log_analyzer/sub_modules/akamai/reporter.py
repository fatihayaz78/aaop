"""DOCX report generator for Akamai analysis using python-docx + kaleido==0.2.1."""

from __future__ import annotations

import io
from pathlib import Path

import structlog
from docx import Document
from docx.shared import Inches

from apps.log_analyzer.config import LogAnalyzerConfig
from apps.log_analyzer.sub_modules.akamai.schemas import AkamaiAnomaly, AkamaiMetrics

logger = structlog.get_logger(__name__)


class AkamaiReporter:
    def __init__(self, config: LogAnalyzerConfig) -> None:
        self._config = config

    def generate(
        self,
        tenant_id: str,
        metrics: AkamaiMetrics,
        anomalies: list[AkamaiAnomaly],
        charts: dict | None = None,
        agent_summary: str | None = None,
    ) -> str:
        """Generate DOCX report and return file path."""
        doc = Document()

        # Cover
        doc.add_heading("Akamai CDN Analysis Report", level=0)
        doc.add_paragraph(f"Tenant: {tenant_id}")
        doc.add_paragraph(f"Total Requests: {metrics.total_requests:,}")
        doc.add_paragraph("")

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        if agent_summary:
            doc.add_paragraph(agent_summary)
        else:
            doc.add_paragraph(
                f"Analysis of {metrics.total_requests:,} requests. "
                f"Error rate: {metrics.error_rate:.2%}, "
                f"Cache hit rate: {metrics.cache_hit_rate:.2%}, "
                f"Avg TTFB: {metrics.avg_ttfb_ms:.1f}ms."
            )

        # Key Metrics Table
        doc.add_heading("Key Metrics", level=1)
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        _set_cell(table, 0, 0, "Metric")
        _set_cell(table, 0, 1, "Value")
        _set_cell(table, 1, 0, "Total Requests")
        _set_cell(table, 1, 1, f"{metrics.total_requests:,}")
        _set_cell(table, 2, 0, "Error Rate")
        _set_cell(table, 2, 1, f"{metrics.error_rate:.2%}")
        _set_cell(table, 3, 0, "Cache Hit Rate")
        _set_cell(table, 3, 1, f"{metrics.cache_hit_rate:.2%}")
        _set_cell(table, 4, 0, "Avg TTFB")
        _set_cell(table, 4, 1, f"{metrics.avg_ttfb_ms:.1f} ms")
        _set_cell(table, 5, 0, "P99 TTFB")
        _set_cell(table, 5, 1, f"{metrics.p99_ttfb_ms:.1f} ms")

        # Anomaly Findings
        doc.add_heading("Anomaly Findings", level=1)
        if anomalies:
            for a in anomalies:
                p = doc.add_paragraph()
                run = p.add_run(f"[{a.severity}] {a.anomaly_type}: ")
                run.bold = True
                p.add_run(a.description)
        else:
            doc.add_paragraph("No anomalies detected.")

        # Charts Gallery
        if charts:
            doc.add_heading("Chart Gallery", level=1)
            for chart_name, fig in charts.items():
                try:
                    img_bytes = fig.to_image(format="png", width=800, height=400)
                    doc.add_heading(chart_name.replace("_", " ").title(), level=2)
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
                except Exception:
                    logger.warning("chart_export_failed", chart=chart_name)
                    doc.add_paragraph(f"[Chart '{chart_name}' could not be exported]")

        # Technical Details
        doc.add_heading("Technical Details", level=1)
        if metrics.top_errors:
            doc.add_heading("Top Errors", level=2)
            err_table = doc.add_table(rows=len(metrics.top_errors) + 1, cols=3)
            err_table.style = "Table Grid"
            _set_cell(err_table, 0, 0, "Error Code")
            _set_cell(err_table, 0, 1, "Count")
            _set_cell(err_table, 0, 2, "Percentage")
            for i, err in enumerate(metrics.top_errors, 1):
                _set_cell(err_table, i, 0, str(err.get("code", "")))
                _set_cell(err_table, i, 1, str(err.get("count", "")))
                _set_cell(err_table, i, 2, f"{err.get('pct', 0):.2f}%")

        # Save
        out_dir = Path(self._config.docx_reports_dir) / tenant_id
        out_dir.mkdir(parents=True, exist_ok=True)
        from datetime import datetime

        filename = f"akamai_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = out_dir / filename
        doc.save(str(path))
        logger.info("docx_report_generated", path=str(path), tenant_id=tenant_id)
        return str(path)


def _set_cell(table: object, row: int, col: int, text: str) -> None:
    table.rows[row].cells[col].text = text  # type: ignore[union-attr]
